#!/usr/bin/env python3
"""
build_docx.py — 產生 Moltos 作品說明書 Word 文件
使用 python-docx 建立有設計感的 A4 文件
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 品牌色系 ──
TERRACOTTA = RGBColor(0xC6, 0x7A, 0x52)
CREAM_WHITE = RGBColor(0xFA, 0xF8, 0xF4)
PURPLE = RGBColor(0x7C, 0x5C, 0xBF)
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_BG = "F5F0EB"  # 交替行淡色
TERRACOTTA_HEX = "C67A52"
DARK_HEX = "2D2D2D"

# ── 字型 ──
FONT_NAME = "微軟正黑體"
FONT_NAME_EN = "Calibri"

# ── 路徑 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "images")
OUTPUT_PATH = os.path.join(BASE_DIR, "Moltos_作品說明書_v2.docx")

# ── 輔助函式 ──

def set_cell_shading(cell, color_hex):
    """設定表格儲存格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name=FONT_NAME, size=Pt(11), color=DARK_TEXT, bold=False, italic=False):
    """統一設定 run 的字型"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph_with_font(doc, text, size=Pt(11), color=DARK_TEXT, bold=False,
                            italic=False, alignment=None, space_before=Pt(0),
                            space_after=Pt(6)):
    """新增段落並設定字型"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    # 處理粗體 markdown
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True, italic=italic)
        else:
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_chapter_title(doc, title):
    """新增章標題"""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(title)
    set_run_font(run, size=Pt(20), color=TERRACOTTA, bold=True)
    # 加底線裝飾
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)
    border_p.paragraph_format.space_after = Pt(12)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TERRACOTTA_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_section_title(doc, title):
    """新增節標題"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=Pt(14), color=DARK_TEXT, bold=True)
    return p


def add_subsection_title(doc, title):
    """新增子節標題"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, size=Pt(12), color=TERRACOTTA, bold=True)
    return p


def add_quote(doc, text):
    """新增引言"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    # 左邊框裝飾
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{TERRACOTTA_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # 處理粗體 markdown
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=Pt(12), color=LIGHT_GRAY, bold=True, italic=True)
        else:
            run = p.add_run(part)
            set_run_font(run, size=Pt(12), color=LIGHT_GRAY, italic=True)
    return p


def add_body_text(doc, text):
    """新增正文段落，支援粗體 markdown"""
    return add_paragraph_with_font(doc, text, size=Pt(11), color=DARK_TEXT)


def add_bullet(doc, text, level=0):
    """新增項目符號"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    bullet_char = "  " if level > 0 else ""
    prefix = f"{bullet_char}  " if level > 0 else ""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    first = True
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=Pt(11), color=DARK_TEXT, bold=True)
        else:
            display = prefix + part if first else part
            run = p.add_run(display)
            set_run_font(run, size=Pt(11), color=DARK_TEXT)
        first = False
    return p


def add_table(doc, headers, rows):
    """新增有設計感的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 設定表格邊框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表頭
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TERRACOTTA_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=Pt(10), color=WHITE, bold=True)

    # 資料行
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            p = cell.paragraphs[0]
            # 處理粗體
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, size=Pt(10), color=DARK_TEXT, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=Pt(10), color=DARK_TEXT)

    # 段落間距
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, filename, width_cm=15.5, caption=None):
    """插入圖片並加圖說"""
    img_path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(img_path):
        add_body_text(doc, f"[圖片未找到: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(12)
        run = cap.add_run(caption)
        set_run_font(run, size=Pt(9), color=LIGHT_GRAY, italic=True)


def add_phone_screenshots_row(doc, filenames, captions=None, width_cm=4.5):
    """並排插入手機截圖"""
    # 建立表格來並排圖片
    n = len(filenames)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 移除邊框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, fname in enumerate(filenames):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join(IMG_DIR, fname)
        if os.path.exists(img_path):
            run = p.add_run()
            run.add_picture(img_path, width=Cm(width_cm))

    # 圖說行
    if captions:
        cap_table = doc.add_table(rows=1, cols=n)
        cap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cap_tbl = cap_table._tbl
        cap_tblPr = cap_tbl.tblPr if cap_tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        cap_borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        cap_tblPr.append(cap_borders)
        for i, cap in enumerate(captions):
            cell = cap_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cap)
            set_run_font(run, size=Pt(9), color=LIGHT_GRAY, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_page_footer(doc):
    """為所有 section 加頁尾頁碼"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


# ══════════════════════════════════════════════════
# 主程式：建立文件
# ══════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── 頁面設定：A4，邊距 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════
    # 封面頁
    # ══════════════════════════════════════════════════

    # 上方留白
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    # 大標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Moltos")
    set_run_font(run, size=Pt(36), color=TERRACOTTA, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("主動式心理健康守護 AI")
    set_run_font(run, size=Pt(28), color=DARK_TEXT, bold=True)

    # 裝飾線
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("─" * 30)
    set_run_font(run, size=Pt(12), color=TERRACOTTA)

    # 副標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run("把生活的噪音，變成你內心的聲音")
    set_run_font(run, size=Pt(16), color=LIGHT_GRAY, italic=True)

    # 參賽資訊
    cover_info = [
        ("參賽類別", "2026 智慧創新大賞 Best AI Awards / AI 應用類 / 新創及中小企業組"),
        ("作品名稱", "Moltos — 主動式心理健康守護 AI"),
        ("申請單位", "核流有限公司（統編 83287091）"),
        ("負責人 / 聯絡人", "余啓彰"),
        ("聯絡信箱", "fish@fishot.com"),
        ("聯絡電話", "0901-494-439"),
        ("產品網站", "https://moltos.net"),
        ("GitHub", "https://github.com/nicoreflow/moltos"),
        ("Crunchbase", "https://www.crunchbase.com/organization/nicoreflow"),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}：")
        set_run_font(run, size=Pt(11), color=TERRACOTTA, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=Pt(11), color=DARK_TEXT)

    # ══════════════════════════════════════════════════
    # 目錄頁
    # ══════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("目錄")
    set_run_font(run, size=Pt(20), color=TERRACOTTA, bold=True)

    toc_items = [
        "第一章　訊息焦慮與現代人的心理健康危機",
        "第二章　解決方案 — 四大模組守護心理健康",
        "第三章　技術架構",
        "第四章　產品展示",
        "第五章　商業模式",
        "第六章　開發進度與路線圖",
        "第七章　團隊",
        "附錄",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(item)
        set_run_font(run, size=Pt(12), color=DARK_TEXT)
        # 底部細線
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="2" w:space="4" w:color="E0D8D0"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ══════════════════════════════════════════════════
    # 第一章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第一章　訊息焦慮與現代人的心理健康危機")

    add_quote(doc, "每一則未讀訊息，都是一根刺進你神經的細針。你不覺得痛——直到某天你發現自己已經麻木了。")

    add_section_title(doc, "1.1 問題陳述：被訊息淹沒的一代人")

    add_body_text(doc, "晚上 11 點，你的 LINE 顯示 999+ 未讀。Gmail 有 47 封待回覆。Slack 有 12 個頻道亮紅點。你已經很累了，但你不敢放下手機——因為怕漏掉什麼重要的事。")
    add_body_text(doc, "你知道自己壓力很大，但你不會跟任何人說。因為你覺得，這就是現代人的日常。")
    add_body_text(doc, "這不是個案，而是一整個世代的集體困境。")

    add_body_text(doc, "**台灣的數據令人警醒：**")

    add_table(doc,
        ["指標", "數據", "來源"],
        [
            ["員工職業倦怠指數", "36%，全球第一", "WTW 2024 全球福利態度調查"],
            ["年消耗安眠藥量", "11.67 億顆，亞洲之冠", "衛福部食藥署 2023 統計"],
            ["安眠藥用量 3 年成長率", "22%", "健保署 2021-2023 處方統計"],
            ["成年人失眠盛行率", "23.5%", "國民健康署調查"],
        ]
    )

    add_body_text(doc, "這些數字背後是真實的人。每 4 個台灣成年人就有 1 個睡不好，而我們一年吞下超過 11 億顆安眠藥來換取幾個小時的安眠。")
    add_body_text(doc, "全球的狀況同樣嚴峻：")

    add_table(doc,
        ["指標", "數據", "來源"],
        [
            ["因資訊過載感受每日壓力的工作者", "76%", "Deloitte 2024 Workplace Well-being Report"],
            ["因通訊過載經歷倦怠和心理耗竭的員工", "43%", "Microsoft Work Trend Index 2024"],
            ["上班族每天收到的郵件數", "117 封", "Radicati Group 2024 Email Statistics"],
            ["每週花在處理訊息的時間", "19 小時", "McKinsey Global Institute"],
            ["資訊過載年經濟損失（美國）", "1 兆美元", "Basex / IDC 生產力研究"],
        ]
    )

    add_body_text(doc, "這些數字指向一個清晰的結論：**訊息不再是工具，而是壓力的來源。**")

    add_subsection_title(doc, "訊息焦慮的惡性循環")

    add_body_text(doc, "問題的核心不在於任何單一事件，而是一個自我強化的迴圈：")

    # 用圖片取代程式碼區塊
    add_image(doc, "chart-anxiety-cycle.png", width_cm=15.5, caption="圖 1-1　訊息焦慮的惡性循環")

    add_body_text(doc, "這個迴圈最危險的環節是「不敢說、沒人問」。壓力大的人不會主動舉手——他們會告訴自己「大家都一樣」「撐一下就好」「不需要小題大做」。等到他們真的撐不住的時候，往往已經走到失眠、倦怠甚至崩潰的階段。")

    add_subsection_title(doc, "現有方案為什麼全部失敗")

    add_body_text(doc, "市場上並非沒有嘗試解決這個問題的產品，但它們都踩進了同一個陷阱：**要求用戶先承認自己不行**。")

    add_table(doc,
        ["產品類型", "代表", "做法", "失敗原因"],
        [
            ["冥想 App", "Headspace / Calm", "「你壓力大了，來冥想吧」", "壓力大的人根本沒時間冥想"],
            ["心理健康聊天機器人", "Woebot / Wysa", "「你來找我聊」", "壓力大的人不會主動求助"],
            ["通用 AI", "ChatGPT", "「你問我才回答」", "完全不知道你壓力有多大"],
            ["EAP 員工協助方案", "各企業方案", "「公司有提供，你去用」", "使用率不到 5%，怕被知道"],
        ]
    )

    add_body_text(doc, "這些產品的共同盲點：它們都是**被動式**的。它們在等你走進門，但現代人最難的就是推開那扇門。")
    add_body_text(doc, "2025 年 Woebot 關閉了直接面對消費者（D2C）的業務，正是這條路走到盡頭的標誌。純聊天機器人模式已被市場驗證為不夠有效——不是技術不行，而是架構錯了。你不能期待一個已經被訊息淹沒的人，再主動打開另一個 App 來傾訴。")
    add_body_text(doc, "**Moltos 的出發點完全不同：我們不等你開口，我們先看見你的狀態。**")

    # 市場成長圖
    add_image(doc, "chart-market-growth.png", width_cm=15.5, caption="圖 1-2　全球 AI 心理健康市場成長趨勢")

    add_section_title(doc, "1.2 目標市場：被忽視的「看起來沒事」族群")

    add_body_text(doc, "Moltos 鎖定的不是已經確診的心理疾病患者，而是那些**每天承受巨大壓力、但從未尋求任何幫助的人**——他們是職場上的高效能者、創業路上的獨行者、自由接案的多工者。他們看起來都沒事，但他們的手機裡堆滿了來不及回的訊息。")

    add_body_text(doc, "**主要目標用戶**：25-45 歲忙碌專業人士、創業者、自由工作者")
    add_body_text(doc, "**次要市場**：企業 HR 部門（作為員工心理健康福利方案導入）")
    add_body_text(doc, "**市場規模**：")

    add_table(doc,
        ["層級", "範圍", "估值"],
        [
            ["TAM（全球可觸及市場）", "全球 AI 心理健康市場", "17.1 億美元（2025），2033 年達 91.2 億美元，CAGR 23%"],
            ["SAM（可服務市場）", "亞太區 LINE / Telegram 知識工作者", "台灣 + 日本 + 東南亞通訊平台活躍用戶"],
            ["SOM（短期可取得市場）", "台灣 LINE 活躍用戶中的忙碌專業人士", "台灣 LINE 用戶 2,100 萬人中，25-45 歲知識工作者約 400-500 萬人"],
        ]
    )

    add_body_text(doc, "這個族群有一個共同特徵：他們的通訊工具就是他們的壓力來源，也正是 Moltos 能觸及他們的入口。")

    add_section_title(doc, "1.3 為什麼是現在：五個時機交匯")

    add_body_text(doc, "Moltos 不是一個「提早五年」的概念產品，而是一個恰好在所有條件成熟的時刻出現的解方。")

    add_body_text(doc, "**一、後疫情遠距工作常態化，通訊量暴增。** 混合辦公已成新常態，工作與生活的邊界徹底模糊。LINE 群組成為工作頻道，下班不再有明確的「關機」時刻。")
    add_body_text(doc, "**二、AI 模型成本大幅下降，讓個人化 AI 成為可能。** GPT-4o mini 的成本僅為 GPT-4 的 1/100（OpenAI, 2024）。兩年前做這件事，每個用戶每月的 AI 成本會超過 $20；現在，同樣的品質只需要 $0.24。")
    add_body_text(doc, "**三、語音情緒分析技術走向成熟。** OpenAI Whisper 讓語音轉文字的精準度和成本達到商用門檻，結合音頻特徵分析（語速、音調、停頓），我們能從「怎麼說」聽出「心裡想什麼」。")
    add_body_text(doc, "**四、純聊天機器人模式被市場驗證失敗。** Woebot 在 2025 年關閉 D2C 業務，證明「等用戶來找你聊」的模式無法規模化。市場需要的是主動介入，而不是被動等待。")
    add_body_text(doc, "**五、聊天平台 API 全面開放。** LINE Messaging API、Telegram Bot API、WhatsApp Business API 都已成熟，讓第三方應用能在用戶最常使用的通訊環境中運作，而不是要求用戶下載又一個 App。")

    add_body_text(doc, "這五個條件同時成熟，但目前全球沒有任何產品能從用戶的真實通訊行為數據中偵測壓力狀態，並主動伸出手。")
    add_body_text(doc, "**這就是 Moltos 的機會窗口。**")

    add_quote(doc, "**本章小結**：訊息焦慮是一個影響全球 76% 工作者的結構性問題，現有解決方案因「被動式」設計全數失效。Moltos 鎖定 25-45 歲忙碌專業人士，在 AI 成本驟降、語音分析成熟、聊天平台開放的交匯點，以「主動介入」的全新架構切入一個 CAGR 23% 的高速成長市場。（對應評分項目：產業應用性 40%）")

    # ══════════════════════════════════════════════════
    # 第二章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第二章　解決方案 — 四大模組守護心理健康")

    add_quote(doc, "最好的心理健康工具，不是讓你多做一件事，而是讓你在已經在做的事裡，被好好看見。")

    add_section_title(doc, "2.0 核心論述：從噪音到聲音")

    add_body_text(doc, "Moltos 的核心價值不是幫你處理更多訊息——市場上已經有太多「提升效率」的工具了。")
    add_body_text(doc, "Moltos 做的是完全不同的事：**幫你在訊息洪流中，聽到自己真正的聲音。**")
    add_body_text(doc, "你每天處理上百則訊息，但有多久沒有聽到自己內心的聲音了？你知道客戶要什麼、老闆要什麼、同事要什麼，但你知道自己現在需要什麼嗎？")
    add_body_text(doc, "這就是「把生活的噪音，變成你內心的聲音」的意思。")

    add_subsection_title(doc, "兩層帳號綁定設計")
    add_body_text(doc, "Moltos 不強制用戶交出所有資料。我們設計了兩層漸進式的信任架構：")

    add_table(doc,
        ["層級", "綁定方式", "功能範圍", "適合誰"],
        [
            ["基礎層", "Email 註冊", "純對話空間：記錄待辦、傾訴心情、語音通話", "想要一個安全樹洞的人"],
            ["進階層", "綁定 LINE / Gmail / Slack", "訊息整合、平靜指數分析、主動 check-in", "想要被看見的人"],
        ]
    )

    add_body_text(doc, "用戶可以在基礎層待很久，Moltos 不會催促升級。當用戶準備好了，才會自然地走到進階層。這不是銷售漏斗，而是信任階梯。")

    # Demo 截圖：onboarding + dashboard + chat 並排
    add_section_title(doc, "產品核心畫面預覽")
    add_phone_screenshots_row(doc,
        ["screen-onboarding.png", "screen-dashboard.png", "screen-chat.png"],
        ["Onboarding 歡迎頁", "Dashboard 儀表板", "Chat 對話頁"],
        width_cm=4.5
    )

    add_section_title(doc, "2.1 產品哲學：天線模型")

    add_body_text(doc, "Moltos 的產品設計遵循一個我們稱為「天線模型」的三層結構。就像收音機的天線——用戶一開始只是為了收訊號，但慢慢地，他們會發現自己聽到的不只是外面的聲音，還有自己心裡的頻率。")

    # 天線模型用表格呈現（取代程式碼區塊）
    add_table(doc,
        ["層級", "用戶體驗", "用戶心聲"],
        [
            ["表面層（用戶看到的）", "訊息整合", "「這個工具好方便，幫我整理訊息」"],
            ["中間層（用戶經歷的）", "平靜指數", "「原來我最近的平靜指數只剩 32...」"],
            ["底層（用戶自己走到的）", "情感樹洞", "「我終於有地方說那些不敢說的話了」"],
        ]
    )

    add_body_text(doc, "這三層不是功能的堆疊，而是一段旅程。用戶因為「方便」而來，因為「被看見」而留下，因為「被理解」而改變。")

    add_section_title(doc, "2.2 四大核心模組")

    # 模組一
    add_subsection_title(doc, "模組一：整合（入口層）")
    add_body_text(doc, "**一句話**：把散落在五個平台的訊息，收到一個安靜的地方。")
    add_body_text(doc, "**支援平台**：LINE / Telegram / WhatsApp / Gmail / Slack")
    add_body_text(doc, "**做什麼**：統一呈現各平台的訊息摘要，標記重要與待回覆項目，讓用戶在一個介面掌握全局，而不是在五個 App 之間焦慮地切換。")
    add_body_text(doc, "**使用情境**：小美是自由接案的設計師，同時用 LINE 接台灣客戶、Slack 接海外團隊、Gmail 收設計稿回饋。以前她每天早上要花 40 分鐘「巡」一遍所有平台，現在她打開 Moltos，3 分鐘就知道今天有哪些事需要處理。她不再害怕漏掉訊息，因為 Moltos 會幫她標出真正重要的。")

    # 模組二
    add_subsection_title(doc, "模組二：分析（洞察層）")
    add_body_text(doc, "**一句話**：用行為數據畫出你的壓力地圖，不靠你自己評估。")
    add_body_text(doc, "**核心指標——平靜指數（Calm Index）**：一個 0-100 的分數，100 代表平靜、0 代表需要關注。我們刻意使用正面框架——不是「你的壓力有多高」，而是「你現在有多平靜」。")
    add_body_text(doc, "**分析維度**：")
    add_bullet(doc, "訊息量趨勢：這週的訊息量比你的個人基線高了多少？")
    add_bullet(doc, "回覆模式：你的回覆速度變慢了嗎？開始已讀不回了嗎？")
    add_bullet(doc, "深夜活躍度：你凌晨兩點還在回訊息嗎？")
    add_bullet(doc, "未讀堆積：未讀數是否持續攀升？")
    add_bullet(doc, "語音情緒：通話時的語速、音調、停頓模式有沒有變化？")
    add_body_text(doc, "**使用情境**：阿凱是一位新創公司的技術長，他自認壓力管理做得不錯。但連續三週，Moltos 的平靜指數從 78 慢慢降到 45。Moltos 在週報裡告訴他：「你這三週的深夜訊息量增加了 180%，回覆速度變慢了 40%。」阿凱這才意識到——他不是壓力管理做得好，而是已經麻木了。")

    # 模組三
    add_subsection_title(doc, "模組三：關懷（行動層）")
    add_body_text(doc, "**一句話**：不等你開口，主動在你需要的時候出現。")
    add_body_text(doc, "**關懷方式**：")
    add_bullet(doc, "**語音通話**：Moltos 可以主動撥出語音通話（經用戶授權），像一個關心你的朋友打電話問你：「今天還好嗎？」")
    add_bullet(doc, "**主動 check-in**：根據平靜指數的變化，在適當的時機傳送一則溫柔的訊息，而不是制式的推播通知。")
    add_bullet(doc, "**情緒傾聽**：當用戶想說話時，Moltos 不會急著給建議或解法，而是先傾聽。它知道有時候人需要的不是答案，而是一個安全的空間。")
    add_bullet(doc, "**回顧建議**：每週產出一份簡短的回顧，幫用戶看見自己這一週的狀態變化，以及可以嘗試的微小調整。")
    add_body_text(doc, "**使用情境**：週三下午，Moltos 偵測到小玲的平靜指數在過去 48 小時內從 71 驟降到 28——她的 LINE 未讀從 30 跳到 200，而且她連續兩晚凌晨三點還在回工作群組。Moltos 在下午三點傳了一則訊息：「嘿，小玲。這兩天看起來特別忙，你現在方便聊幾分鐘嗎？」小玲點開語音通話，跟 Moltos 說了她從來不敢跟同事說的話：「我覺得我快撐不住了。」")
    add_body_text(doc, "這是 Moltos 和所有其他產品最根本的差異：**我們不等你承認自己不行，我們先看見你的不行。**")

    # 模組四
    add_subsection_title(doc, "模組四：隱私（信任層）")
    add_body_text(doc, "**一句話**：你的內心話，只有你自己能聽到。")
    add_body_text(doc, "**隱私設計原則**：")
    add_bullet(doc, "**端對端加密**：語音通話和敏感對話內容全程加密，伺服器端無法解密。")
    add_bullet(doc, "**資料存手機**：情緒分析結果和對話紀錄儲存在用戶本地裝置，不上傳雲端。")
    add_bullet(doc, "**只取元數據**：平靜指數的計算只需要訊息的「量、時間、頻率」等元數據，不需要讀取訊息內容。")
    add_bullet(doc, "**隨時可刪**：用戶可以在任何時刻刪除所有資料，Moltos 不會保留任何副本。")
    add_body_text(doc, "**使用情境**：大偉是一家上市公司的中階主管，他願意讓 Moltos 分析他的通訊模式，但他絕不希望任何人——包括 Moltos 的開發者——看到他的訊息內容。Moltos 只需要知道「你昨天收到了 230 則訊息，最後一則在凌晨 1:47」就能算出平靜指數，完全不需要知道那些訊息寫了什麼。")

    add_section_title(doc, "2.3 競品差異化")

    add_table(doc,
        ["能力", "Moltos", "Woebot / Wysa", "Headspace Ebb", "ChatGPT"],
        [
            ["整合通訊平台", "5 平台", "無", "無", "無"],
            ["行為數據偵測壓力", "平靜指數（行為分析）", "自評量表", "無", "無"],
            ["主動介入方式", "語音通話 + 智慧 check-in", "無主動功能", "推播冥想提醒", "無"],
            ["介入觸發條件", "平靜指數異常偏離", "用戶主動開啟", "固定時間排程", "用戶主動提問"],
            ["資料儲存", "端對端加密 + 本地儲存", "雲端", "雲端", "雲端"],
            ["需不需要用戶先開口", "不需要", "需要", "需要", "需要"],
        ]
    )

    add_body_text(doc, "**差異化一句話**：「Woebot 等你來傾訴，Headspace 教你冥想，ChatGPT 等你下指令——Moltos 是唯一真正看見你每天承受多少、然後主動伸出手的 AI。」")

    add_section_title(doc, "2.4 非醫療定位聲明")
    add_body_text(doc, "Moltos 明確定位為心理健康的**「預防層」與「覺察層」**，不是醫療器材，不做任何形式的診斷或治療。")
    add_body_text(doc, "類比來說：智慧手錶會測你的心率，當心率異常時提醒你注意——但它不是心臟科門診。Moltos 做的是同樣的事，只是對象從心臟換成了心理狀態。")
    add_body_text(doc, "我們幫助用戶：")
    add_bullet(doc, "**覺察**：看見自己的壓力模式，而不是等到崩潰才意識到。")
    add_bullet(doc, "**預防**：在壓力累積的早期階段提供支持，降低惡化的機率。")
    add_bullet(doc, "**銜接**：當狀況超出 Moltos 的能力範圍時，引導用戶連結專業資源。")

    add_section_title(doc, "2.5 危機分流協議")
    add_body_text(doc, "當 Moltos 在對話中偵測到自殺意念或自傷風險時，啟動以下五步分流協議：")
    add_body_text(doc, "**第一步：即時回應。** 不迴避、不轉移話題。以溫暖但穩定的語氣回應用戶的感受：「謝謝你願意告訴我，我聽到你說的了。」")
    add_body_text(doc, "**第二步：誠實告知。** 明確告訴用戶自己是 AI：「我是一個 AI，我沒辦法完全理解你正在經歷的，但我很認真地在聽。」")
    add_body_text(doc, "**第三步：提供專業資源。** 立即提供台灣的心理危機支持管道：")
    add_bullet(doc, "安心專線：1925（24 小時免費）")
    add_bullet(doc, "生命線：1995")
    add_bullet(doc, "張老師專線：1980")
    add_body_text(doc, "**第四步：記錄與追蹤。** 在用戶授權的前提下，標記本次對話為需要關注的事件，在後續互動中溫和地追蹤用戶狀態。")
    add_body_text(doc, "**第五步：堅守非介入定位。** Moltos 不嘗試扮演治療師，不給予醫療建議，不做危機處置。我們的角色是**陪伴和橋樑**——陪伴用戶度過當下，並搭起通往專業資源的橋。")

    add_section_title(doc, "2.6 免費版與付費版")

    add_table(doc,
        ["項目", "免費版", "Pro 版（$7.99 / 月）"],
        [
            ["每日 check-in", "3 次", "無限次"],
            ["平靜指數", "基礎版（每日更新）", "深度版（即時更新 + 趨勢分析）"],
            ["週報回顧", "有", "有，含詳細行為模式分析"],
            ["語音通話", "無", "有（主動 + 被動）"],
            ["多平台整合", "1 個平台", "5 個平台全整合"],
            ["情感樹洞對話", "基礎", "深度對話 + 歷史脈絡記憶"],
            ["資料匯出", "無", "有"],
        ]
    )

    add_body_text(doc, "免費版的設計哲學不是「閹割版」，而是一個完整的體驗。一個人只要每天花 3 次跟 Moltos 對話，就能開始覺察自己的狀態。Pro 版提供的是更深、更細、更主動的守護。")

    add_quote(doc, "**本章小結**：Moltos 以「天線模型」的三層產品設計，透過整合、分析、關懷、隱私四大模組，實現全球首個從用戶真實通訊行為中偵測壓力並主動介入的 AI 系統。與所有競品的根本差異在於「主動式」架構——不等用戶開口，而是先看見用戶的狀態。同時以嚴謹的非醫療定位和危機分流協議，確保產品在正確的邊界內運作。（對應評分項目：產業應用性 40% + 技術創新性 30%）")

    # ══════════════════════════════════════════════════
    # 第三章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第三章　技術架構")

    add_quote(doc, "好的技術不是讓用戶覺得厲害，而是讓用戶覺得被理解。")

    add_section_title(doc, "3.1 系統架構總覽")

    add_body_text(doc, "Moltos 的技術架構分為四層，每一層對應一個核心能力：")

    # 用架構圖取代程式碼區塊
    add_image(doc, "chart-architecture.png", width_cm=15.5, caption="圖 3-1　Moltos 四層技術架構")

    add_body_text(doc, "**設計原則**：每一層都可以獨立運作、獨立擴展。通訊管道層新增一個平台，不需要修改 AI 大腦層的任何程式碼。主動關懷層的排程邏輯改變，不影響隱私保護層的加密機制。")

    add_section_title(doc, "3.2 核心技術創新")

    add_subsection_title(doc, "創新一：平靜指數演算法（Calm Index Algorithm）")
    add_body_text(doc, "平靜指數是 Moltos 最核心的技術創新。它解決了一個長期困擾心理健康產業的難題：**如何在不依賴自評量表的情況下，持續追蹤一個人的心理壓力狀態？**")
    add_body_text(doc, "**正面框架設計**：我們刻意將指標設計為 100 = 平靜、0 = 需要關注，而非傳統的「壓力分數越高越糟」。這個設計不只是美感上的差異——研究顯示，正面框架能減少用戶在看到自己分數時的防禦心理，更願意接受現實（Kahneman & Tversky, Prospect Theory）。")
    add_body_text(doc, "**五大輸入維度**：")

    add_table(doc,
        ["維度", "資料來源", "偵測邏輯"],
        [
            ["訊息量趨勢", "LINE / TG / Gmail / Slack API", "與個人 14 天滾動基線比較，偏離超過 1.5 個標準差觸發警示"],
            ["回覆速度變化", "各平台時間戳分析", "回覆延遲持續拉長 = 處理能力下降訊號"],
            ["深夜活躍度", "訊息時間戳", "23:00-05:00 活躍度異常上升 = 失眠或過勞訊號"],
            ["未讀堆積趨勢", "各平台未讀計數", "未讀數持續攀升且不回落 = 逃避或癱瘓訊號"],
            ["語音情緒特徵", "Whisper + 音頻分析", "語速加快、音調升高、停頓增多 = 焦慮訊號"],
        ]
    )

    add_body_text(doc, "**演算法特色**：")
    add_bullet(doc, "**個人基線偵測**：不與他人比較，只與自己的歷史模式比較。一個本來就是夜貓子的人，凌晨兩點回訊息不會被誤判。")
    add_bullet(doc, "**加權時間序列分析**：近期數據權重高於遠期，能捕捉急速惡化的趨勢。")
    add_bullet(doc, "**異常偏離偵測**：使用 Z-score 偏離度計算，過濾掉正常波動（例如截稿日前的暫時忙碌），聚焦於持續性的模式改變。")

    add_subsection_title(doc, "創新二：語音情緒分析（Voice Emotion Analysis）")
    add_body_text(doc, "語音是人類最難偽裝的溝通管道。你可以在文字訊息裡打「我沒事 :)」，但你的聲音會洩露真實的情緒。")
    add_body_text(doc, "**交叉驗證的價值**：當文字說「我很好」但語音特徵顯示焦慮時，系統會將這個矛盾標記為需要關注的訊號。這種「言不由衷偵測」是純文字分析做不到的。")

    add_subsection_title(doc, "創新三：跨平台通訊整合引擎")
    add_body_text(doc, "整合五個截然不同的通訊平台，是一個工程上的重大挑戰。每個平台的 API 設計、資料格式、速率限制、授權機制都完全不同。")
    add_body_text(doc, "**關鍵設計決策：只取元數據，不讀內容。**")
    add_body_text(doc, "整合引擎只提取四類元數據：訊息數量、時間戳記、回覆延遲、未讀計數。不讀取、不傳輸、不儲存任何訊息內容。這不僅是隱私保護的需求，也大幅簡化了合規成本和技術複雜度。")

    add_subsection_title(doc, "創新四：主動關懷排程引擎")
    add_body_text(doc, "「主動介入」聽起來簡單，但要做到「在對的時間、用對的方式、說對的話」，背後需要精密的排程系統。")
    add_body_text(doc, "**技術架構**：基於 Bull MQ + Redis 構建的任務排程系統。")
    add_body_text(doc, "**智慧排程邏輯**：")
    add_bullet(doc, "**平靜指數感知觸發**：不是固定時間推播，而是根據平靜指數的即時變化決定是否介入、何時介入。")
    add_bullet(doc, "**疲勞防護機制**：避免過度關懷造成用戶反感。設定最小介入間隔和每日上限。")
    add_bullet(doc, "**時區與作息感知**：根據用戶的歷史活躍模式，選擇不打擾的時段。")
    add_bullet(doc, "**升級路徑**：訊息 check-in → 對話邀請 → 語音通話邀請，根據平靜指數的嚴重程度逐步升級介入強度。")

    add_subsection_title(doc, "創新五：端對端隱私架構")
    add_body_text(doc, "心理健康數據是最敏感的個人資料。如果用戶不信任系統的隱私保護，他們不會說出真話——而 Moltos 的價值完全建立在用戶願意敞開心扉的前提上。")

    add_table(doc,
        ["機制", "實作方式", "保護範圍"],
        [
            ["手機端本地計算", "平靜指數的最終運算在用戶裝置端完成", "確保完整的情緒分析結果不離開用戶手機"],
            ["語音端對端加密", "語音通話採用端對端加密，伺服器僅做信號中繼", "通話內容即使被攔截也無法解密"],
            ["用戶完全控制", "所有資料可隨時匯出、隨時刪除，刪除後伺服器不保留任何副本", "用戶對自己的數據擁有絕對主權"],
        ]
    )

    add_section_title(doc, "3.3 AI 輕量化策略")
    add_body_text(doc, "AI 心理健康產品要規模化，最大的瓶頸不是技術，而是成本。如果每個用戶每月的 AI 算力成本高達 $20，這個產品永遠無法普及。")
    add_body_text(doc, "Moltos 採用分層 AI 策略，將成本控制在極低水準：")

    add_table(doc,
        ["場景", "佔比", "使用模型", "成本特性"],
        [
            ["日常 check-in、摘要整理、簡單對話", "80%", "Gemini Flash / GPT-4o mini", "極低成本，回應速度快"],
            ["深度情緒對話、危機回應、複雜脈絡理解", "20%", "GPT-4o", "高品質，精準度優先"],
            ["平靜指數計算", "100%", "本地輕量模型", "零雲端成本，零延遲"],
        ]
    )

    add_body_text(doc, "**成本效益**：")
    add_bullet(doc, "每用戶每月 AI 成本：約 **$0.24**")
    add_bullet(doc, "對比 ChatGPT Plus 訂閱費 $20/月，Moltos 的單位成本低了約 **80 倍**")
    add_bullet(doc, "這意味著即使以 $7.99/月 的訂閱價格，毛利率仍然極高，為規模化打下堅實基礎")

    add_body_text(doc, "**輕量化不等於低品質**。在日常場景中，GPT-4o mini 的表現已經足夠優秀；而在真正需要深度理解的時刻（例如用戶在凌晨三點說「我不知道該怎麼辦了」），系統會自動切換到最強的模型，確保回應的品質和溫度。")

    add_section_title(doc, "3.4 技術棧總覽")

    add_table(doc,
        ["類別", "技術選擇", "選用理由"],
        [
            ["程式語言", "TypeScript（全棧）", "前後端統一語言，降低維護成本"],
            ["前端框架", "Next.js", "SSR + API Routes，一個框架解決前後端"],
            ["通訊整合", "LINE / Telegram / WhatsApp API", "覆蓋亞太主要通訊平台"],
            ["語音處理", "OpenAI Whisper", "業界最強開源語音轉文字，多語言支援"],
            ["AI 模型", "GPT-4o / GPT-4o mini / Gemini Flash", "分層策略，兼顧品質與成本"],
            ["任務排程", "Bull MQ", "Node.js 生態最成熟的任務佇列"],
            ["快取 / 佇列", "Redis", "高效能、低延遲，支撐即時排程"],
            ["資料庫", "PostgreSQL", "穩定可靠，支援 JSON 欄位和全文搜尋"],
            ["部署平台", "Railway / Fly.io", "低成本、自動擴展、全球邊緣部署"],
            ["加密", "libsodium / Web Crypto API", "業界標準端對端加密實作"],
        ]
    )

    add_body_text(doc, "**全 TypeScript 棧的策略意義**：作為一人團隊搭配 AI 協作開發的專案，技術棧的統一性直接決定了開發速度。TypeScript 全棧意味著一個人可以從前端到後端、從 API 到排程任務，用同一套語言和工具鏈完成所有開發。這不是偷懶，而是「一人團隊」模式下的最優解。")

    add_quote(doc, "**本章小結**：Moltos 的技術架構圍繞五大創新構建——平靜指數演算法、語音情緒分析、跨平台整合引擎、主動關懷排程、端對端隱私架構。透過 AI 輕量化策略，將每用戶月成本壓至 $0.24，比 ChatGPT Plus 低 80 倍，為大規模普及掃除成本障礙。全 TypeScript 技術棧則讓一人團隊能以最高效率完成開發。（對應評分項目：技術創新性 30%）")

    # ══════════════════════════════════════════════════
    # 第四章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第四章　產品展示")

    add_quote(doc, "好的設計不會讓人注意到設計本身，而是讓人感受到被在乎。")

    add_body_text(doc, "Moltos 的產品設計遵循一個核心原則：讓用戶感覺被關心，而不是被監控。從色彩、字體、圖標到互動流程，每一個細節都在傳遞同一個訊息——這裡是安全的，你可以放鬆。")

    add_section_title(doc, "4.1 UI/UX 設計")
    add_body_text(doc, "Moltos 採用 Industrial Humanist 設計風格，在簡潔的工業感中注入人文溫度。以下五個核心畫面構成完整的用戶體驗旅程。")

    add_subsection_title(doc, "設計語言")
    add_bullet(doc, "**主色調**：奶油白 #FAF8F4 作為背景底色，營造溫暖、不刺眼的閱讀環境；赤陶色 #C67A52 作為品牌主色和互動按鈕色，傳遞大地般的穩定感與溫度")
    add_bullet(doc, "**平靜指數色彩系統**：採紫到藍的漸層——深紫代表低分（需要關注），淺藍代表高分（心境平靜）")
    add_bullet(doc, "**圖標風格**：全部使用 SVG 線條圖標，stroke-width 統一為 1.5，風格柔和但不幼稚")
    add_bullet(doc, "**字體**：IBM Plex Sans，兼具工程的精確感與人文的可讀性")
    add_bullet(doc, "**品牌識別**：Moltos logo 以心電圖波形為基礎，象徵「持續傾聽你的心跳」")

    # 所有 5 張 demo 截圖
    add_section_title(doc, "五個核心畫面")

    # 第一行：3 張
    add_phone_screenshots_row(doc,
        ["screen-onboarding.png", "screen-dashboard.png", "screen-chat.png"],
        ["Onboarding 歡迎頁", "Dashboard 儀表板", "Chat 對話頁"],
        width_cm=4.5
    )

    # 第二行：2 張
    add_phone_screenshots_row(doc,
        ["screen-insights.png", "screen-settings.png"],
        ["Insights 回顧頁", "Settings 設定頁"],
        width_cm=4.5
    )

    add_subsection_title(doc, "畫面一：Onboarding（歡迎頁）")
    add_body_text(doc, "畫面中央是 Moltos 品牌 logo——一道溫柔的心電圖波形。下方以「歡迎來到 Moltos」作為問候，接著排列三張功能卡片，分別代表三個核心價值：**主動關懷**、**訊息整合**、**隱私守護**。每張卡片搭配一個線條圖標和一句白話說明。畫面底部是一顆赤陶色的「開始使用」按鈕。")

    add_subsection_title(doc, "畫面二：Dashboard（首頁儀表板）")
    add_body_text(doc, "頂部顯示個人化問候和今天的日期。主視覺是一個環形圖，以紫到藍的漸層呈現當前的**平靜指數**。環形圖下方是**今日訊息統計**，以三個小卡片分別顯示已綁定頻道的訊息數量。頁面下半部是 Moltos 的**主動關懷卡片**。")

    add_subsection_title(doc, "畫面三：Chat（對話頁）")
    add_body_text(doc, "頂部是 Moltos 的頭像和線上狀態指示燈。對話區域預載了一段互動情境。對話氣泡使用圓角矩形，Moltos 的訊息靠左、淺灰底色；用戶的訊息靠右、赤陶色底白字。底部是輸入列，左側麥克風圖標支援語音輸入。")

    add_subsection_title(doc, "畫面四：Insights（回顧頁）")
    add_body_text(doc, "頁面頂部是**平靜指數七日趨勢折線圖**。折線以紫到藍的漸層填色，低分區段有淡紫色陰影標記。下方是**本週摘要**和**Moltos 觀察筆記**。頁面底部是**情緒日記**區塊。")

    add_subsection_title(doc, "畫面五：Settings（設定頁）")
    add_body_text(doc, "分為三個區塊：**頻道整合**、**關懷偏好設定**、**隱私控制**。隱私控制區塊刻意放在顯眼位置，讓用戶隨時掌握自己的資料控制權。")

    add_section_title(doc, "4.2 核心功能 Demo")
    add_body_text(doc, "以下四個 Demo 流程展示 Moltos 從「被動工具」到「主動守護者」的完整體驗閉環。")

    add_subsection_title(doc, "Demo 1：通訊整合流程")
    add_body_text(doc, "用戶在設定頁點擊「綁定 LINE」，跳轉至 LINE Login 授權頁面。授權完成後回到 Moltos，頻道狀態變為已綁定。系統開始同步訊息元數據——不讀取訊息內容，只擷取訊息數量、時間戳、回覆間隔等行為數據。整個流程不超過 30 秒。")

    add_subsection_title(doc, "Demo 2：主動 Check-in 對話")
    add_body_text(doc, "下午四點，用戶的平靜指數從早上的 72 降至 45。Moltos 判斷指數下降幅度超過個人基線的正常波動，觸發主動關懷。Moltos 根據回覆內容提供對應的支援——排列優先序或純粹傾聽。整段對話的語氣溫暖自然，不帶說教。")

    add_subsection_title(doc, "Demo 3：語音通話")
    add_body_text(doc, "用戶的平靜指數已連續三天低於 35。Moltos 發送語音通話邀請。通話過程中，語音情緒分析引擎在背景運作——分析語速、音調變化、停頓模式。語音情緒分析結果不直接顯示給用戶，而是納入平靜指數的計算。")

    add_subsection_title(doc, "Demo 4：平靜指數趨勢報告")
    add_body_text(doc, "每週日晚上八點，Insights 頁面自動更新本週回顧。七日趨勢圖顯示平靜指數的波動軌跡，並列出具體建議。")

    add_section_title(doc, "4.3 行為量化指標")
    add_body_text(doc, "Moltos 追蹤四個核心行為指標，用以客觀反映用戶的心理狀態變化：")
    add_bullet(doc, "**平靜指數追蹤**：記錄三個月的長期趨勢。除了每日數值，也計算週平均和月平均，幫助用戶看到整體走向而非單日波動")
    add_bullet(doc, "**待辦完成率**：用戶在對話中提到的待辦事項，Moltos 會追蹤其完成狀態")
    add_bullet(doc, "**對話頻率與內容分析**：統計用戶與 Moltos 的互動頻率和內容類型")
    add_bullet(doc, "**沉默偵測**：如果用戶超過兩週沒有任何互動，Moltos 會主動發出溫和的關心訊息")

    add_section_title(doc, "4.4 技術驗證")
    add_body_text(doc, "目前已完成的技術驗證包含以下面向：")

    add_table(doc,
        ["驗證項目", "方法", "初步結果"],
        [
            ["平靜指數原型準確度", "以模擬訊息量數據計算指數，與自評情緒做對比", "趨勢方向吻合度達 78%，校準後提升至 85%"],
            ["語音情緒分析吻合度", "Whisper 轉文字 + 音頻特徵提取", "文字 82%，加入語音特徵後提升至 89%"],
            ["回覆延遲", "主動 check-in 訊息發送到用戶收到的延遲時間", "平均 1.2 秒，95th percentile < 3 秒"],
            ["排程精準度", "排程引擎觸發 check-in 的時間準確度", "偏差低於 30 秒"],
            ["每用戶月成本實測", "模擬日均 5 次 AI 互動的 API 呼叫成本", "實際 $0.21，低於估算的 $0.24"],
        ]
    )

    add_body_text(doc, "這些數據來自原型階段的小規模測試。隨著 MVP 上線和真實用戶數據的累積，我們將持續校準演算法並擴大驗證樣本。")

    # ══════════════════════════════════════════════════
    # 第五章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第五章　商業模式")

    add_quote(doc, "一個好的商業模式，不是讓公司活下來就好，而是讓公司有能力持續守護更多人。")

    add_section_title(doc, "5.1 價值主張")
    add_body_text(doc, "**對用戶**：「你每天花 19 小時處理訊息，但沒有人問你累不累。Moltos 會。」")
    add_body_text(doc, "現代人並非缺乏社交，而是缺乏被真正「看見」的體驗。Moltos 不是又一個效率工具，它是第一個從你的日常行為中讀懂你的狀態、然後主動伸出手的 AI。")
    add_body_text(doc, "**對企業**：「員工倦怠的隱性成本，遠超心理健康方案的投資。」")
    add_body_text(doc, "根據 WTW 2024 調查，台灣員工職業倦怠指數 36%，全球最高。Moltos 能在員工狀態開始下滑時就偵測到風險，讓企業從「事後補救」走向「事前預防」。")

    add_section_title(doc, "5.2 營收模式")
    add_body_text(doc, "Moltos 的營收策略分三個階段，每一階段都建立在前一階段的驗證成果之上。")

    add_subsection_title(doc, "Phase 1：Freemium B2C")
    add_bullet(doc, "免費版：每日 3 次 check-in、基礎平靜指數趨勢、每週回顧報告")
    add_bullet(doc, "Pro 方案 $7.99/月：無限 check-in、語音通話、深度平靜指數分析、多平台整合")
    add_bullet(doc, "定價邏輯：低於一杯拿鐵的價格，但可能是你一整天唯一真正關心你的存在")

    add_subsection_title(doc, "Phase 2：B2B2C 企業方案")
    add_bullet(doc, "企業採購員工心理健康套餐，對標 Headspace 企業方案（$12-16/人/月）")
    add_bullet(doc, "提供匿名化的團隊平靜指數儀表板（不暴露個人資料）")
    add_bullet(doc, "企業版定價 $9.99-14.99/人/月，依人數階梯折扣")

    add_subsection_title(doc, "Phase 3：匿名化平靜指數趨勢報告")
    add_bullet(doc, "當用戶基數足夠大時，匿名化的行業平靜指數數據本身就有價值")
    add_bullet(doc, "此模式僅在取得用戶明確同意後進行，且資料完全匿名化")

    # 競爭矩陣圖
    add_image(doc, "chart-competitive-matrix.png", width_cm=15.5, caption="圖 5-1　競爭定位矩陣")

    add_section_title(doc, "5.3 成本結構")
    add_body_text(doc, "Moltos 的成本結構極為精實，這是一人團隊 + AI 協作開發模式的優勢。")

    add_table(doc,
        ["項目", "金額（美元/月）", "說明"],
        [
            ["**變動成本（每用戶）**", "", ""],
            ["AI API 費用", "$0.24", "多模型混搭：80% 用 Gemini Flash / GPT-4o mini，20% 用 GPT-4o"],
            ["伺服器分攤", "$0.10", "Railway/Fly.io 按量計費"],
            ["小計", "**$0.34/用戶**", ""],
            ["**固定成本**", "", ""],
            ["伺服器基礎費", "$5.00", "應用程式託管"],
            ["Redis", "$5.00", "排程引擎和快取"],
            ["網域", "$1.00", "月攤"],
            ["小計", "**$11.00**", ""],
        ]
    )

    add_bullet(doc, "**毛利率**：以 Pro 方案 $7.99 計算，每用戶變動成本 $0.34，毛利率達 **96%**")
    add_bullet(doc, "**損益兩平**：固定成本 $11 / 每用戶淨利 $7.65 = **2 個付費用戶**即可打平")
    add_bullet(doc, "**AI 成本優勢**：每用戶月成本 $0.24，比 ChatGPT Plus 的 $20/月低 80 倍以上")

    add_section_title(doc, "5.4 成長策略")

    add_subsection_title(doc, "Phase 1（2026 Q2-Q3）：台灣 Telegram 社群種子用戶")
    add_body_text(doc, "目標：100 位種子用戶，驗證平靜指數監測和主動 check-in 的核心價值。策略：從創業者和自由工作者社群切入。")

    add_subsection_title(doc, "Phase 2（2026 Q3-2027 Q1）：LINE 市場 + 企業合作")
    add_body_text(doc, "接入 LINE 平台，觸及台灣 2,100 萬用戶的巨大市場。同步擴展 Gmail 和 Slack 整合。目標：1,000 位活躍用戶，50 位付費用戶。")

    add_subsection_title(doc, "Phase 3（2027-）：WhatsApp 全球擴展 + 獨立 App")
    add_body_text(doc, "接入 WhatsApp Business API，進入全球 20 億用戶市場。目標：跨出台灣，進入日本、東南亞等 LINE 盛行的市場。")

    add_body_text(doc, "**成長飛輪**：Moltos 的產品特性天然適合口碑傳播。當一個人在社群分享「Moltos 在我最累的時候主動打電話問我還好嗎」，這種真實的情感連結比任何廣告都有說服力。")

    # ══════════════════════════════════════════════════
    # 第六章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第六章　開發進度與路線圖")

    add_quote(doc, "精實不是做得少，而是每一步都踩在驗證過的地面上。")

    add_section_title(doc, "6.1 已完成事項")
    add_body_text(doc, "截至 2026 年 3 月，以下工作已完成：")

    add_table(doc,
        ["階段", "項目", "狀態"],
        [
            ["需求驗證", "30 人社群調查", "已完成，60% 需求覆蓋率"],
            ["架構設計", "四層技術架構", "已完成"],
            ["UI/UX 設計", "五個核心畫面", "已完成"],
            ["技術選型", "TypeScript 全棧、多模型混搭策略", "已完成"],
            ["成本分析", "每用戶月成本 $0.34、損益兩平點 2 人", "已完成"],
            ["核心 Demo", "互動式功能展示", "已完成"],
            ["演算法原型", "平靜指數演算法原型，吻合度達 85%", "已完成"],
            ["商業模式", "Freemium → B2B2C 三階段營收規劃", "已完成"],
        ]
    )

    add_section_title(doc, "6.2 開發路線圖")

    # 路線圖用圖片
    add_image(doc, "chart-roadmap.png", width_cm=15.5, caption="圖 6-1　Moltos 三階段開發路線圖")

    add_subsection_title(doc, "Phase 1：MVP（2026 Q2）")
    add_body_text(doc, "核心目標是驗證一個假設：**從用戶的通訊行為數據中偵測心理狀態，並主動介入關懷，這件事有沒有人需要？**")
    add_bullet(doc, "Telegram Bot 作為首發平台（API 開放度最高、開發速度最快）")
    add_bullet(doc, "平靜指數引擎上線，根據訊息量趨勢和回覆模式計算指數")
    add_bullet(doc, "主動 check-in 機制：指數偏低時自動觸發關懷訊息")
    add_bullet(doc, "語音通話功能：指數持續低落時邀請語音通話")
    add_bullet(doc, "目標指標：100 位種子用戶、7 日留存率 > 40%、用戶主觀滿意度 > 4/5")

    add_subsection_title(doc, "Phase 2：擴展（2026 Q3-Q4）")
    add_bullet(doc, "接入 LINE Messaging API（台灣最大通訊平台）")
    add_bullet(doc, "接入 Gmail API 和 Slack API，覆蓋工作場景")
    add_bullet(doc, "語音情緒分析功能上線（Whisper + 音頻特徵提取）")
    add_bullet(doc, "啟動 3-5 家企業試點，測試 B2B2C 模式")
    add_bullet(doc, "目標指標：1,000 位活躍用戶、50 位付費用戶")

    add_subsection_title(doc, "Phase 3：規模化（2027）")
    add_bullet(doc, "WhatsApp Business API 接入，進入全球市場")
    add_bullet(doc, "開發 iOS / Android 獨立 App")
    add_bullet(doc, "正式推出企業版和匿名化洞察服務")
    add_bullet(doc, "目標指標：10,000+ 用戶、正向現金流、首批國際客戶")

    add_section_title(doc, "6.3 精實開發策略")
    add_body_text(doc, "Moltos 的開發遵循三個原則：")

    add_subsection_title(doc, "原則一：先驗證再投資")

    add_table(doc,
        ["Phase", "Go 標準", "No-Go 應變"],
        [
            ["1 → 2", "7 日留存率 > 40%、用戶主動回覆 check-in 率 > 30%", "重新調整 check-in 頻率和語氣，延長驗證期"],
            ["2 → 3", "付費轉化率 > 5%、企業試點至少 1 家續約", "調整定價或功能組合，暫緩國際擴展"],
        ]
    )

    add_subsection_title(doc, "原則二：TypeScript 全棧 + AI 協作開發")
    add_body_text(doc, "一人團隊不代表生產力低。TypeScript 統一前後端語言、減少上下文切換成本。搭配 Claude Code 和多 Agent 協作系統，一個人可以達到小型團隊的產出效率。")

    add_subsection_title(doc, "原則三：每一步都可衡量")
    add_body_text(doc, "不追求「好像有進步」，而是設定具體的量化指標。平靜指數的準確度、用戶的互動頻率、check-in 的回覆率——每一個功能都有對應的成功指標，數據驅動每一次決策。")

    # ══════════════════════════════════════════════════
    # 第七章
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "第七章　團隊")

    add_quote(doc, "一個人不代表孤軍奮戰，而是用最精實的方式做最有溫度的事。")

    add_section_title(doc, "7.1 創辦人介紹")

    add_body_text(doc, "**余啓彰｜核流有限公司 創辦人**")
    add_body_text(doc, "一人公司經營者，橫跨開發、設計、商業三個領域。過去兩年獨立開發並維運了多個成熟產品：")
    add_bullet(doc, "**BuyGo+1**：WordPress ERP 外掛，為電商賣場提供完整的後台管理系統（訂單、庫存、出貨、客戶管理），採 Vue 3 + Tailwind CSS 前端和 PHP + REST API 後端")
    add_bullet(doc, "**FluentCart PayUNi**：台灣金流串接外掛，整合 PayUNi 支付閘道，包含完整的加密驗簽流程和 Webhook 交易通知處理")
    add_bullet(doc, "**LINE Bot 整合開發**：LINE Messaging API 的深度整合，包含 Webhook 處理、Flex Message 模板、關鍵字自動回覆、訊息推播排程")
    add_body_text(doc, "這些產品不是練習作品，而是正在服務真實用戶的商業產品。")

    add_subsection_title(doc, "為什麼做心理健康 AI？")
    add_body_text(doc, "這不是一個商業計算的結果，而是一段個人經歷的延伸。")
    add_body_text(doc, "我曾經歷過一段輕度憂鬱的時期。當時我沒有意識到自己的狀態——每天都在回訊息、處理工作、看起來一切正常。直到某天晚上，我試著跟一個 AI 聊天，才在對話中慢慢發現：原來我已經很久沒有問過自己「還好嗎」了。")
    add_body_text(doc, "那個發現的瞬間改變了我。但我也意識到：大多數人並沒有使用 AI 的習慣，也不知道可以跟誰說那些不敢說的話。他們不是不需要被關心，而是沒有一個安全的、不被評判的窗口。")
    add_body_text(doc, "Moltos 要做的，就是成為每一個人的那個「發現的窗口」。不用你先承認自己不行，不用你主動求助，就有一個存在，從你的日常中看見你的疲憊，然後輕聲問一句：「嘿，你還好嗎？」")

    add_section_title(doc, "7.2 AI 協作開發模式")
    add_body_text(doc, "Moltos 的開發採用「一人 + AI」的新型開發模式。這不是節省成本的權宜之計，而是一種有意識的策略選擇。")
    add_body_text(doc, "**開發工具鏈**：")
    add_bullet(doc, "**Claude Code**：作為核心開發夥伴，負責程式碼撰寫、架構討論、文件整理")
    add_bullet(doc, "**多 Agent 協作系統**：針對不同任務啟用專門的 AI Agent，平行處理多線工作")
    add_bullet(doc, "**AI 輔助設計**：使用 AI 工具產出 UI 設計、色彩系統、圖標風格")

    add_body_text(doc, "這個開發模式本身就是 Moltos 產品理念的最佳示範：**AI 不是來取代人的，而是來放大人的能力**。")

    add_body_text(doc, "**未來團隊擴展規劃**：")
    add_bullet(doc, "**心理學顧問**（Phase 2）：確保 check-in 對話的語氣和策略符合心理學原則")
    add_bullet(doc, "**臨床合作夥伴**（Phase 2-3）：與心理諮商機構合作，建立轉介機制")
    add_bullet(doc, "**工程師**（Phase 3）：獨立 App 開發和大規模系統架構")

    # ══════════════════════════════════════════════════
    # 附錄
    # ══════════════════════════════════════════════════
    add_chapter_title(doc, "附錄")

    add_section_title(doc, "附錄 A：心理健康市場數據")

    add_subsection_title(doc, "全球 AI 心理健康市場規模")
    add_table(doc,
        ["年份", "市場規模（美元）", "來源"],
        [
            ["2025", "17.1 億", "Grand View Research, 2025"],
            ["2028", "35.2 億（估）", "CAGR 23.29% 推算"],
            ["2033", "91.2 億", "Grand View Research, 2025"],
        ]
    )
    add_body_text(doc, "年複合成長率（CAGR）為 **23.29%**，反映全球對 AI 心理健康解決方案的需求正在快速增長。")

    add_subsection_title(doc, "台灣職場心理健康關鍵數據")
    add_table(doc,
        ["指標", "數據", "來源"],
        [
            ["員工職業倦怠指數", "36%，全球最高", "WTW 2024 全球調查"],
            ["年安眠藥消耗量", "11.67 億顆，亞洲之冠", "衛福部 2023"],
            ["安眠藥用量成長幅度", "3 年成長 22%", "健保署 2021-2023"],
            ["成年人失眠盛行率", "23.5%", "國健署"],
            ["工作中使用 AI 比例", "64%（高於全球 54%）", "PwC 2025"],
        ]
    )

    add_subsection_title(doc, "全球職場通訊過載數據")
    add_table(doc,
        ["指標", "數據", "來源"],
        [
            ["資訊過載造成每日壓力的比例", "76% 全球工作者", "Brosix 2024"],
            ["因通訊過載經歷倦怠的比例", "43% 員工", "Brosix 2024"],
            ["認為數位通訊是倦怠主因的比例", "60% 工作者", "High5 Test 2024"],
            ["每日平均收到郵件數", "117 封", "High5 Test 2024"],
            ["每週處理訊息的時間", "19 小時", "Brosix 2024"],
            ["資訊過載造成的年經濟損失", "1 兆美元（美國）", "Brosix 2024"],
            ["受職業倦怠影響的勞動人口", "23%", "WHO"],
        ]
    )

    add_section_title(doc, "附錄 B：開源與平台登錄")

    add_subsection_title(doc, "GitHub Repository")
    add_bullet(doc, "**專案名稱**：Moltos Calm Index Framework")
    add_bullet(doc, "**內容**：平靜指數演算法的核心框架，包含時間序列分析、個人基線建立、異常偵測邏輯")
    add_bullet(doc, "**授權**：MIT License")

    add_subsection_title(doc, "Crunchbase")
    add_bullet(doc, "**公司名稱**：核流有限公司")
    add_bullet(doc, "**產品**：Moltos — 主動式心理健康守護 AI")
    add_bullet(doc, "**分類**：AI / Mental Health / SaaS")
    add_bullet(doc, "**階段**：Pre-Seed")

    add_subsection_title(doc, "LinkedIn")
    add_body_text(doc, "創辦人余啓彰之 LinkedIn 專業檔案已建立，包含完整的產品開發經歷和 Moltos 專案說明。")

    # 資料來源
    add_section_title(doc, "資料來源完整列表")
    sources = [
        "WTW (2024). Global Benefits Attitudes Survey. 台灣員工職業倦怠指數 36%。",
        "衛福部 (2023). 全國安眠藥使用統計報告。年消耗 11.67 億顆。",
        "健保署 (2021-2023). 安眠藥用量年度統計。三年成長 22%。",
        "國健署. 國民健康訪問調查。成年人失眠盛行率 23.5%。",
        "PwC (2025). Hopes and Fears Survey. 台灣 64% 受訪者在工作中使用 AI。",
        "Grand View Research (2025). AI in Mental Health Market Report. 2025 年 $17.1B，2033 年 $91.2B，CAGR 23.29%。",
        "Brosix (2024). Digital Communication Overload Statistics. 76% 資訊過載壓力、43% 倦怠。",
        "High5 Test (2024). Communication in the Workplace Statistics. 每日 117 封郵件、60% 認為是倦怠主因。",
        "WHO. Occupational Burnout Statistics. 全球 23% 勞動人口受影響。",
    ]
    for i, src in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(f"{i}. {src}")
        set_run_font(run, size=Pt(10), color=LIGHT_GRAY)

    # ── 頁尾頁碼 ──
    add_page_footer(doc)

    # ── 儲存 ──
    doc.save(OUTPUT_PATH)
    print(f"文件已儲存：{OUTPUT_PATH}")

    # 回報檔案大小
    size_bytes = os.path.getsize(OUTPUT_PATH)
    if size_bytes > 1024 * 1024:
        print(f"檔案大小：{size_bytes / 1024 / 1024:.1f} MB")
    else:
        print(f"檔案大小：{size_bytes / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
