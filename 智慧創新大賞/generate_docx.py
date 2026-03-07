#!/usr/bin/env python3
"""產生 Moltos 作品說明書（5頁精簡版）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ===== 版面設定：A4 直向，邊界上下左右 2cm，裝訂邊 1cm =====
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(3.0)   # 左 2cm + 裝訂邊 1cm
section.right_margin = Cm(2.0)

# ===== 預設段落樣式：14pt 標楷體、單行間距、前後段距 3pt =====
style = doc.styles['Normal']
font = style.font
font.name = '標楷體'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
pf = style.paragraph_format
pf.line_spacing = 1.0
pf.space_before = Pt(3)
pf.space_after = Pt(3)

# ===== 標題樣式 =====
for i in [1, 2, 3]:
    hs = doc.styles[f'Heading {i}']
    hf = hs.font
    hf.name = '標楷體'
    hf.bold = True
    hf.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    hp = hs.paragraph_format
    hp.line_spacing = 1.0
    hp.space_before = Pt(6)
    hp.space_after = Pt(3)

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(14)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '標楷體'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(14)
    if bold:
        run.bold = True
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    return p


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = '標楷體'
        r1.font.size = Pt(13)
        r1.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        r2 = p.add_run(text)
        r2.font.name = '標楷體'
        r2.font.size = Pt(13)
        r2.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    else:
        r = p.add_run(text)
        r.font.name = '標楷體'
        r.font.size = Pt(13)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    pf2 = p.paragraph_format
    pf2.line_spacing = 1.0
    pf2.space_before = Pt(1)
    pf2.space_after = Pt(1)
    return p


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = '標楷體'
        r.font.size = Pt(12)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        if bold:
            r.bold = True
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)


def make_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = '標楷體'
        r.font.size = Pt(12)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
    for row_data in rows:
        add_table_row(table, row_data)
    return table


# =====================================================================
# 一、創意發想背景及概述
# =====================================================================
add_heading('一、創意發想背景及概述', level=1)

add_para(
    '現代人每天處理超過 117 封 Email、數百則即時訊息，76% 的工作者因資訊過載感受每日壓力'
    '（Deloitte, 2024）。台灣更是全球職業倦怠指數第一（36%，WTW 2024），年消耗安眠藥 11.67 '
    '億顆（衛福部 2023），每 4 位成年人就有 1 位失眠。訊息焦慮已形成「訊息爆量→焦慮累積→'
    '壓力內化→效率下降→更多訊息」的惡性循環。',
    indent=True
)

add_para(
    '然而，現有心理健康產品（Woebot、Wysa、Headspace、ChatGPT）全部採用被動式架構——'
    '等用戶主動求助。2025 年 Woebot 關閉消費者業務，證明純聊天機器人模式已走到盡頭。'
    '壓力最大的人，恰恰是最不會主動開口的人。',
    indent=True
)

add_para(
    'Moltos（主動式心理健康守護 AI）的核心創新在於：不等用戶開口，從其真實通訊行為數據'
    '中偵測壓力狀態，並主動伸出援手。我們的產品標語是「把生活的噪音，變成你內心的聲音」。'
    'Moltos 是全球首個整合多平台通訊元數據、即時計算心理狀態指標、並主動發起語音關懷的 '
    'AI 系統。',
    indent=True
)

# =====================================================================
# 二、作品功能簡介及特色
# =====================================================================
add_heading('二、作品功能簡介及特色', level=1)

add_para('Moltos 以四大核心模組構成完整的守護閉環：', bold=True)

add_bullet(
    '統一呈現 LINE、Telegram、WhatsApp、Gmail、Slack 五大平台的訊息摘要，標記重要與待回覆項目，'
    '用戶在一個介面掌握全局，不再在多個 App 間焦慮切換。',
    bold_prefix='模組一｜訊息整合：'
)

add_bullet(
    '獨創「平靜指數」（Calm Index, 0-100），從訊息量趨勢、回覆速度變化、深夜活躍度、未讀堆積、'
    '語音情緒五大維度，以個人 14 天滾動基線進行 EWMA + Z-score 異常偵測，客觀量化壓力狀態，'
    '無需用戶自評。',
    bold_prefix='模組二｜行為分析：'
)

add_bullet(
    '當平靜指數異常下降時，Moltos 主動發起文字 check-in 或語音通話邀請。語音通話同步進行'
    '雙通道情緒分析（語義內容 + 音頻特徵），在用戶最需要的時刻主動伸出手。每週自動產出'
    '趨勢回顧與建議。',
    bold_prefix='模組三｜主動關懷：'
)

add_bullet(
    '平靜指數在用戶端本地計算；語音通話端對端加密；僅擷取元數據（訊息數量、時間戳），'
    '絕不讀取訊息內容；用戶可隨時匯出或刪除所有資料。',
    bold_prefix='模組四｜隱私守護：'
)

add_para('關鍵特色：', bold=True)

add_bullet(
    '全球唯一從真實通訊行為數據偵測壓力並主動介入的 AI，「主動式」架構根本區隔於所有被動式競品。',
    bold_prefix='主動式架構：'
)

add_bullet(
    '採正面框架（100=平靜、0=需關注），減少用戶防禦心理，演算法已開源（MIT License）。',
    bold_prefix='平靜指數演算法：'
)

add_bullet(
    '明確定位為心理健康「預防層」與「覺察層」，非醫療器材。偵測到自殺意念時啟動五步危機分流'
    '協議，引導連結專業資源（安心專線 1925 等）。',
    bold_prefix='非醫療定位：'
)

# =====================================================================
# 三、開發工具與技術
# =====================================================================
add_heading('三、開發工具與技術', level=1)

add_heading('3.1 技術架構', level=2)

add_para(
    'Moltos 採用四層分離架構：通訊管道層（五大平台 API 接入）→ AI 大腦層（平靜指數演算法、'
    '語音情緒分析、對話引擎）→ 主動關懷層（Bull MQ 排程引擎、智慧觸發）→ 隱私保護層'
    '（端對端加密、本地儲存、元數據分離）。各層獨立運作、獨立擴展。',
    indent=True
)

add_heading('3.2 技術棧', level=2)

make_table(
    ['類別', '技術選擇', '說明'],
    [
        ['程式語言', 'TypeScript 全棧', '前後端統一語言，一人團隊最優解'],
        ['前端框架', 'Next.js', 'SSR + API Routes，一框架解決前後端'],
        ['通訊整合', 'LINE / Telegram / WhatsApp API', '覆蓋亞太主要通訊平台'],
        ['語音處理', 'OpenAI Whisper', '業界最強開源語音轉文字'],
        ['AI 模型', 'GPT-4o / GPT-4o mini / Gemini Flash', '多模型混搭策略'],
        ['排程引擎', 'Bull MQ + Redis', '任務佇列與即時排程'],
        ['資料庫', 'PostgreSQL', '穩定可靠，支援 JSON 與全文搜尋'],
        ['加密', 'libsodium / Web Crypto API', '端對端加密實作'],
    ]
)

add_heading('3.3 技術獨特性與差異化', level=2)

add_bullet(
    '以 EWMA 建立個人 14 天滾動基線，Z-score 偵測偏離 > 1.5 標準差的異常，'
    'Sigmoid 評分 + 跨維度交叉懲罰，實現「只跟自己比」的個人化壓力量化。已開源（MIT License）。',
    bold_prefix='平靜指數演算法：'
)

add_bullet(
    'Whisper 語音轉文字（語義分析）+ 音頻特徵提取（語速、音調、停頓、顫抖），'
    '雙通道交叉驗證可偵測「言不由衷」，為純文字分析不可能做到的能力。',
    bold_prefix='語音情緒雙通道分析：'
)

add_bullet(
    '80% 日常互動用輕量模型（Gemini Flash / GPT-4o mini），20% 深度對話用 GPT-4o。'
    '每用戶月成本僅 $0.24，比 ChatGPT Plus 低約 80 倍，為大規模普及掃除成本障礙。',
    bold_prefix='AI 輕量化策略：'
)

add_heading('3.4 迭代與擴充規劃', level=2)

make_table(
    ['階段', '時程', '內容'],
    [
        ['Phase 1 MVP', '2026 Q2', 'Telegram Bot + 平靜指數 + 主動 check-in + 語音通話'],
        ['Phase 2 擴展', '2026 Q3-Q4', '接入 LINE / Gmail / Slack、語音情緒分析上線、企業試點'],
        ['Phase 3 規模化', '2027', 'WhatsApp 接入、iOS/Android App、B2B 企業方案、國際市場'],
    ]
)

# =====================================================================
# 四、使用對象及環境
# =====================================================================
add_heading('四、使用對象及環境', level=1)

add_heading('4.1 使用對象', level=2)

add_bullet(
    '25-45 歲忙碌專業人士、創業者、自由工作者——每天被訊息淹沒但不會主動求助的族群。',
    bold_prefix='主要目標：'
)
add_bullet(
    '企業 HR 部門，作為員工心理健康福利方案導入，提供匿名化團隊平靜指數儀表板。',
    bold_prefix='次要市場：'
)

add_heading('4.2 軟硬體環境', level=2)

make_table(
    ['項目', '規格說明'],
    [
        ['用戶端', '任何可運行 LINE / Telegram / WhatsApp 的智慧型手機或電腦（iOS / Android / Web）'],
        ['通訊平台', 'LINE Messaging API、Telegram Bot API、WhatsApp Business API、Gmail API、Slack API'],
        ['伺服器', 'Railway / Fly.io 雲端部署，自動擴展，全球邊緣節點'],
        ['AI 運算', '雲端：OpenAI API + Google Gemini API；端側：平靜指數本地計算（零雲端成本）'],
        ['資料庫', 'PostgreSQL（結構化資料）+ Redis（快取與排程）'],
        ['加密', '端對端加密（libsodium），語音通話加密中繼，資料本地儲存'],
        ['未來端側 AI', '規劃整合 Qualcomm AI Engine / Hexagon DSP，實現語音情緒分析端側推論'],
    ]
)

# =====================================================================
# 五、產業應用性
# =====================================================================
add_heading('五、產業應用性', level=1)

add_heading('5.1 市場規模與成長性', level=2)

add_para(
    '全球 AI 心理健康市場 2025 年達 17.1 億美元，預計 2033 年成長至 91.2 億美元（CAGR 23.29%'
    '，Grand View Research）。台灣 LINE 用戶 2,100 萬人中，25-45 歲知識工作者約 400-500 萬人'
    '為 Moltos 的可觸及市場。',
    indent=True
)

add_heading('5.2 商業價值', level=2)

add_bullet(
    'Pro 方案 $7.99/月，每用戶變動成本僅 $0.34/月，毛利率達 96%。',
    bold_prefix='高毛利模式：'
)
add_bullet(
    '固定成本 $11/月，2 位付費用戶即可損益兩平，不需大量融資即可存活。',
    bold_prefix='極低損益點：'
)
add_bullet(
    'B2C Freemium → B2B2C 企業方案（$9.99-14.99/人/月）→ 匿名化產業洞察報告。',
    bold_prefix='三階段營收：'
)

add_heading('5.3 社會影響力', level=2)

add_para(
    'Moltos 鎖定心理健康的「預防層」，在壓力累積早期即主動介入，有望降低因長期壓力導致的'
    '失眠、倦怠與心理疾病發生率。台灣每年 11.67 億顆安眠藥的消耗背後，是數百萬人未被'
    '看見的壓力。Moltos 的主動式架構讓「不敢開口的人」也能被關懷到，填補傳統心理健康'
    '服務最大的缺口。平靜指數演算法以 MIT License 開源，促進學術界與產業界共同推進'
    '行為數據驅動的心理健康研究。',
    indent=True
)

add_heading('5.4 預計合作對象', level=2)

add_bullet(
    '整合 Qualcomm AI Engine / Hexagon DSP，將語音情緒分析與平靜指數運算移至端側，'
    '實現零延遲、零雲端成本、隱私最大化的 always-on 監測。',
    bold_prefix='Qualcomm：'
)
add_bullet(
    '接入 LINE Messaging API 觸及台灣 2,100 萬用戶，並擴展至日本、東南亞市場。',
    bold_prefix='LINE / 通訊平台：'
)
add_bullet(
    '心理諮商機構、企業 HR 部門、保險業者，建立轉介機制與企業員工心理健康福利方案。',
    bold_prefix='醫療與企業：'
)

# =====================================================================
# 六、結語
# =====================================================================
add_heading('六、結語', level=1)

add_para(
    'Moltos 不是又一個等你開口的聊天機器人，而是全球首個從真實通訊行為中看見你的壓力、'
    '並主動伸出手的 AI 守護者。在 AI 模型成本驟降、語音分析技術成熟、通訊平台 API 全面'
    '開放的時機交匯點，Moltos 以獨創的平靜指數演算法、主動式關懷架構與端對端隱私設計，'
    '切入一個 CAGR 23% 的高速成長市場。我們相信：最好的心理健康守護，不是讓人多做一件事，'
    '而是在他已經在做的事裡，被好好看見。',
    indent=True
)

# ===== 頁碼 =====
from docx.oxml import OxmlElement
sectPr = section._sectPr
pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:start'), '1')
sectPr.append(pgNumType)

# Footer with page number
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar1)

run2 = fp.add_run()
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' PAGE '
run2._r.append(instrText)

run3 = fp.add_run()
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run3._r.append(fldChar2)

for r in [run, run2, run3]:
    r.font.name = '標楷體'
    r.font.size = Pt(12)

# ===== 儲存 =====
output_path = '/Users/fishtv/Development/care-engine/智慧創新大賞/上傳資料/附件三_作品說明書.docx'
doc.save(output_path)
print(f'已儲存至 {output_path}')

# 也存一份到主目錄
import shutil
shutil.copy(output_path, '/Users/fishtv/Development/care-engine/智慧創新大賞/Moltos_作品說明書_v3_5頁.docx')
print('副本已存至 智慧創新大賞/Moltos_作品說明書_v3_5頁.docx')
